import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER

from app.schemas.drafting import StructuredDocumentObject

class DocumentGenerator:
    @staticmethod
    def generate_docx(doc_obj: StructuredDocumentObject) -> io.BytesIO:
        doc = Document()
        
        # Set margins to 1 inch
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        if doc_obj.title:
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_p.add_run(doc_obj.title.replace('₹', 'Rs. '))
            run.font.underline = True
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            doc.add_paragraph() # spacing

        # Parties
        for key, value in doc_obj.parties.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{key.replace('_', ' ').title()}: ")
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run_val = p.add_run(value.replace('₹', 'Rs. '))
            run_val.font.name = 'Times New Roman'
            run_val.font.size = Pt(12)
            
        doc.add_paragraph()

        # Body
        for para in doc_obj.body:
            p = doc.add_paragraph(para.replace('₹', 'Rs. '))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        doc.add_paragraph()

        # Verification
        if doc_obj.verification and doc_obj.verification.text:
            p = doc.add_paragraph("VERIFICATION")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            
            p_text = doc.add_paragraph(doc_obj.verification.text)
            p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p_text.runs:
                run.font.name = 'Times New Roman'
            
            p_date = doc.add_paragraph(f"Date: {doc_obj.verification.date}")
            p_date.runs[0].font.name = 'Times New Roman'
            p_place = doc.add_paragraph(f"Place: {doc_obj.verification.place}")
            p_place.runs[0].font.name = 'Times New Roman'
            
        # Signature Blocks
        if doc_obj.signature_blocks:
            doc.add_paragraph("\n\n")
            p = doc.add_paragraph()
            for sig in doc_obj.signature_blocks:
                run = p.add_run(f"______________________\n{sig}\n\n")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Annexures
        if doc_obj.annexures:
            doc.add_page_break()
            p = doc.add_paragraph("ANNEXURES")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            for idx, ann in enumerate(doc_obj.annexures, 1):
                p = doc.add_paragraph(f"{idx}. {ann}")
                p.runs[0].font.name = 'Times New Roman'

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    @staticmethod
    def generate_pdf(doc_obj: StructuredDocumentObject) -> io.BytesIO:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=72)
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('TitleStyle', parent=styles['Normal'], fontName='Times-Bold', fontSize=14, alignment=TA_CENTER, spaceAfter=20)
        body_style = ParagraphStyle('BodyStyle', parent=styles['Normal'], fontName='Times-Roman', fontSize=12, alignment=TA_JUSTIFY, spaceAfter=12)
        bold_style = ParagraphStyle('BoldStyle', parent=styles['Normal'], fontName='Times-Bold', fontSize=12, spaceAfter=6)
        right_style = ParagraphStyle('RightStyle', parent=styles['Normal'], fontName='Times-Roman', fontSize=12, alignment=2, spaceAfter=30)
        
        story = []
        
        if doc_obj.title:
            sanitized_title = doc_obj.title.replace('₹', 'Rs. ')
            story.append(Paragraph(f"<u>{sanitized_title}</u>", title_style))
            
        for key, value in doc_obj.parties.items():
            sanitized_val = value.replace('₹', 'Rs. ')
            story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {sanitized_val}", body_style))
            
        story.append(Spacer(1, 12))
        
        for para in doc_obj.body:
            sanitized_para = para.replace('₹', 'Rs. ')
            
            # Simple markdown bold parsing for ReportLab
            if '**' in sanitized_para:
                parts = sanitized_para.split('**')
                for i in range(1, len(parts), 2):
                    parts[i] = f"<b>{parts[i]}</b>"
                sanitized_para = "".join(parts)
                
            story.append(Paragraph(sanitized_para, body_style))
            
        story.append(Spacer(1, 12))
        
        if doc_obj.verification and doc_obj.verification.text:
            story.append(Paragraph("VERIFICATION", title_style))
            story.append(Paragraph(doc_obj.verification.text, body_style))
            story.append(Paragraph(f"Date: {doc_obj.verification.date}", body_style))
            story.append(Paragraph(f"Place: {doc_obj.verification.place}", body_style))
            
        if doc_obj.signature_blocks:
            story.append(Spacer(1, 40))
            for sig in doc_obj.signature_blocks:
                story.append(Paragraph("______________________", right_style))
                story.append(Paragraph(sig, right_style))
                
        if doc_obj.annexures:
            # ReportLab doesn't easily do a manual page break here without a PageBreak object
            from reportlab.platypus import PageBreak
            story.append(PageBreak())
            story.append(Paragraph("ANNEXURES", title_style))
            for idx, ann in enumerate(doc_obj.annexures, 1):
                story.append(Paragraph(f"{idx}. {ann}", body_style))

        doc.build(story)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_reasoning_pdf(content: str) -> io.BytesIO:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        styles = getSampleStyleSheet()
        
        # very basic markdown parsing for reportlab
        story = []
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 12))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], styles['Heading2']))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], styles['Heading1']))
            elif line.startswith('- '):
                story.append(Paragraph(f"• {line[2:]}", styles['BodyText']))
            elif line.startswith('**') and line.endswith('**'):
                story.append(Paragraph(f"<b>{line[2:-2]}</b>", styles['BodyText']))
            else:
                # Replace inline bold
                line = line.replace('**', '<b>', 1).replace('**', '</b>', 1)
                story.append(Paragraph(line, styles['BodyText']))
                
        doc.build(story)
        buffer.seek(0)
        return buffer
