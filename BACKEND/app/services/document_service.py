import os
import uuid
import shutil
import fitz # PyMuPDF
import docx
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from app.models.document import Document
from app.schemas.upload_chat import DocumentUploadResponse
from app.core.config import settings
from google import genai
from google.genai import types

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class DocumentService:
    def __init__(self):
        api_key = settings.GEMINI_API_KEY or "DUMMY_KEY_FOR_TESTING"
        self.client = genai.Client(api_key=api_key)

    def process_upload(self, user_uid: str, file: UploadFile, db: Session) -> DocumentUploadResponse:
        doc_id = str(uuid.uuid4())
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in [".pdf", ".docx"]:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF or DOCX.")
            
        filepath = os.path.join(UPLOAD_DIR, f"{doc_id}{ext}")
        with open(filepath, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        # Check size limit (10MB)
        file_size = os.path.getsize(filepath)
        if file_size > 10 * 1024 * 1024:
            os.remove(filepath)
            raise HTTPException(status_code=400, detail="File size exceeds the 10MB limit.")
            
        text = ""
        pages = 0
        
        try:
            if ext == ".pdf":
                doc = fitz.open(filepath)
                pages = len(doc)
                if pages > 300:
                    doc.close()
                    os.remove(filepath)
                    raise HTTPException(status_code=400, detail="Document exceeds the maximum limit of 300 pages.")
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            elif ext == ".docx":
                doc = docx.Document(filepath)
                pages = 1 # Approximation for DOCX
                for para in doc.paragraphs:
                    text += para.text + "\n"
        except HTTPException:
            raise
        except Exception as e:
            if os.path.exists(filepath):
                os.remove(filepath)
            raise HTTPException(status_code=400, detail=f"Failed to process document: {str(e)}")
            
        if not text.strip():
            os.remove(filepath)
            raise HTTPException(status_code=400, detail="Document is empty or text could not be extracted.")
            
        # Generate Summary
        summary = self._generate_summary(text)
        
        db_doc = Document(
            id=doc_id,
            user_uid=user_uid,
            filename=file.filename,
            filepath=filepath,
            pages=pages,
            extracted_text=text,
            summary=summary
        )
        db.add(db_doc)
        db.commit()
        
        return DocumentUploadResponse(
            document_id=doc_id,
            filename=file.filename,
            pages=pages,
            summary=summary
        )

    def _generate_summary(self, text: str) -> str:
        prompt = "Summarize the following legal document in 2-3 concise sentences. Focus on the nature of the document and its primary purpose:\n\n" + text[:15000]
        
        import time
        import logging
        import concurrent.futures
        logger = logging.getLogger(__name__)
        
        max_retries = 4
        models_to_try = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-flash-latest', 'gemini-flash-lite-latest']
        
        for attempt in range(max_retries):
            model_name = models_to_try[attempt % len(models_to_try)]
            try:
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(
                        self.client.models.generate_content,
                        model=model_name,
                        contents=prompt
                    )
                    response = future.result(timeout=15)
                return response.text.strip()
            except concurrent.futures.TimeoutError as e:
                logger.warning(f"Summary generation with {model_name} timed out after 15 seconds.")
                if attempt < max_retries - 1:
                    continue
            except Exception as e:
                logger.warning(f"Summary generation with {model_name} failed: {str(e)}")
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)
                    continue
        
        return "This document was uploaded successfully, but an AI summary is temporarily unavailable. You can still ask questions about the document."

document_service = DocumentService()
